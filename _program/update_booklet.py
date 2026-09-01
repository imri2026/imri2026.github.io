#!/usr/bin/env python3
"""
update_booklet.py — Rebuild the timetable and author index in the Word
program booklet from _program/program.csv, _program/posters.csv, and
_program/authors_accepted.csv.

Usage:
    python update_booklet.py [source.docx] [output.docx]

Defaults:
    source: iMRI2026-program-booklet.docx   (same directory as this script)
    output: iMRI2026-program-booklet-updated.docx

The timetable (the first table in the document) is rebuilt. The last
table in the document — originally "Regular Session Speakers" — is
replaced with a full "Author Index": every author (not just presenters)
of every accepted abstract, alphabetical by last name, with a fourth
column listing every O-#/P-# presentation number they're an author on.
The second-to-last table (originally "Invited Speakers (2024 meeting)")
is replaced with "Invited Speakers": every invited_talk/keynote row in
program.csv that has an assigned speaker, alphabetical by last name.
Everything else in the document (cover text, styles) is left untouched.

Each session becomes its own small table (caption row + Start/End/Title
header + one row per talk), mirroring the per-session tables on the
website's program page, instead of one continuous table for the whole
day. Non-session rows (breaks, welcome, reception, etc.) are plain
paragraphs between the tables. Row shading/fonts are cloned from the
original document's template rows, so formatting matches the source
exactly; only text content changes.
"""

import copy
import csv
import re
import sys
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path

import docx
from docx.shared import RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

NAVY = RGBColor(0x25, 0x47, 0x7A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Column widths (twips): Start | End | # (gold accent) | Title/Presentation.
# Widened from the original 360-twip decorative strip so "O-16"/"P-95" fit.
COL_WIDTHS = (900, 810, 720, 8370)
TOTAL_WIDTH = sum(COL_WIDTHS)

# Author-index column widths (twips): Name | Affiliation | City, Country |
# Presentations. Sums to the same 10710 total as the original 3-column
# "Regular Session Speakers" table so the table stays full-width.
AUTHOR_COL_WIDTHS = (2000, 3600, 2400, 2710)

SCRIPT_DIR = Path(__file__).parent

CHILD_TYPES = ('invited_talk', 'oral', 'keynote')


def load_csv(path):
    with open(path, newline='', encoding='utf-8') as f:
        return [{k: v.strip() for k, v in row.items()} for row in csv.DictReader(f)]


def clean_text(s):
    """Collapse embedded newlines/whitespace from source spreadsheet cells
    into single spaces, so multi-line source data renders as one clean
    line in the docx table cell."""
    return re.sub(r'\s+', ' ', (s or '')).strip()


def last_name_key(name):
    """Sort key for alphabetizing a roster by last name."""
    parts = name.split()
    last = parts[-1].lower() if parts else ''
    return (last, name.lower())


CREDENTIAL_RE = re.compile(
    r'\s*,?\s*\b(PhD|Ph\.D\.|MD|M\.D\.|MSc|BSc|DSc|RN|DrPH|FACR|FSIR|FCIRSE|VMD|Dr\.?|Prof\.?)\.?\s*$',
    re.IGNORECASE,
)


def strip_credentials(name):
    """Strip trailing credentials (PhD, MD, ...) from a name, repeatedly
    (some names carry more than one, e.g. "Jane Doe MD PhD")."""
    prev = None
    while prev != name:
        prev = name
        name = CREDENTIAL_RE.sub('', name).strip()
    return name


def build_subid_to_pres(program_rows, poster_rows):
    """Sub_ID -> ('O', number) / ('P', number) for every accepted abstract
    that made it into the final program, from program.csv's subid column
    (oral talks) and posters.csv's subid column (posters)."""
    mapping = {}
    for r in program_rows:
        if r.get('type') == 'oral' and r.get('subid') and r.get('number'):
            mapping[r['subid']] = ('O', int(r['number']))
    for r in poster_rows:
        if r.get('subid') and r.get('number'):
            mapping[r['subid']] = ('P', int(r['number']))
    return mapping


def build_author_index(authors_csv, subid_to_pres):
    """Load authors_accepted.csv (one row per author per submission) and
    collapse it into one row per unique person, each carrying every
    presentation number (O-#/P-#) they're an author on — as presenter or
    co-author — sorted alphabetically by last name.

    People are grouped by email (the reliable identity key); the handful
    of rows with no email fall back to grouping by exact name. A second
    pass then merges any groups that still share the same normalized name
    (e.g. someone who submitted under two different institutional emails)
    into one person.
    """
    rows = load_csv(authors_csv)

    def group_key(r):
        email = r.get('Email', '').strip().lower()
        return ('email', email) if email else ('name', r.get('Author', '').strip().lower())

    groups = defaultdict(list)
    for r in rows:
        groups[group_key(r)].append(r)

    def norm_name(n):
        n = re.sub(r'[.\-]', ' ', n)
        return re.sub(r'\s+', ' ', n).strip().lower()

    merged = defaultdict(list)
    for k, rs in groups.items():
        display_name = Counter(r['Author'].strip() for r in rs).most_common(1)[0][0]
        merged[norm_name(display_name)].extend(rs)

    authors = []
    for _, rs in merged.items():
        name = Counter(r['Author'].strip() for r in rs).most_common(1)[0][0]
        affiliation = Counter(clean_text(r['Affiliation']) for r in rs).most_common(1)[0][0]
        city = Counter(clean_text(r['City']) for r in rs).most_common(1)[0][0]
        country = Counter(clean_text(r['Country']) for r in rs).most_common(1)[0][0]
        location = ', '.join(filter(None, [city, country]))

        subids = set(r['Sub_ID'] for r in rs)
        pres = sorted({subid_to_pres[s] for s in subids if s in subid_to_pres})
        pres_str = ', '.join(f'{kind}-{num}' for kind, num in pres)

        authors.append({
            'name': clean_text(name),
            'affiliation': affiliation,
            'location': location,
            'presentations': pres_str,
        })

    authors.sort(key=lambda a: last_name_key(a['name']))
    return authors


def build_invited_speakers(program_rows):
    """Every invited_talk/keynote row in program.csv that has an assigned
    speaker (i.e. not a TBA placeholder or a note-only row), collapsed to
    Name | Affiliation | City, Country and sorted alphabetically by last
    name. The `speaker` column mixes two conventions in the source data —
    "Name CREDENTIAL, Affiliation" all in one field, or a bare name with
    the affiliation in its own column — both are handled."""
    speakers = []
    for r in program_rows:
        if r.get('type') not in ('invited_talk', 'keynote'):
            continue
        raw = (r.get('speaker') or '').strip()
        if not raw:
            continue

        affiliation = (r.get('affiliation') or '').strip()
        if affiliation:
            name_part = raw
        elif ',' in raw:
            name_part, affiliation = raw.split(',', 1)
        else:
            name_part = raw

        name = strip_credentials(name_part.strip())
        location = ', '.join(filter(None, [(r.get('city') or '').strip(), (r.get('country') or '').strip()]))

        speakers.append({
            'name': name,
            'affiliation': clean_text(affiliation),
            'location': location,
        })

    speakers.sort(key=lambda s: last_name_key(s['name']))
    return speakers


def fmt_time(t):
    return t or ''


def fmt_day_banner(date_str):
    d = datetime.strptime(date_str, '%Y-%m-%d')
    return f'{d.strftime("%A, %B")} {d.day}, {d.year}'.upper()


def status_text(s):
    s = (s or '').lower()
    if s == 'invited':
        return ' (Invited)'
    if s in ('tba', 'tbd'):
        return ' (TBD)'
    return ''


def speaker_line(row):
    speaker = row.get('speaker', '') or ''
    affil = row.get('affiliation', '') or ''
    city = row.get('city', '') or ''
    country = row.get('country', '') or ''
    if not speaker:
        return None
    line = speaker
    if affil:
        line += f', {affil}'
    location = ', '.join(filter(None, [city, country]))
    if location:
        line += f', {location}'
    line += status_text(row.get('status', ''))
    return line


# ---------------------------------------------------------------------------
# docx building blocks
# ---------------------------------------------------------------------------

def set_run(para, text, bold=None):
    """Set a paragraph's text using its existing first run (so its cloned
    font/size/color are preserved) instead of python-docx's default
    text-replacement, which resets formatting."""
    runs = para.runs
    if runs:
        run = runs[0]
        run.text = text
        for extra in runs[1:]:
            extra._element.getparent().remove(extra._element)
    else:
        run = para.add_run(text)
    if bold is not None:
        run.bold = bold
    return run


def set_styled_run(cell, text, color, bold=True):
    """Like set_run, but for cells whose template paragraph has no run at
    all (e.g. the gold accent column, never previously used for text) —
    creates one with explicit formatting instead of relying on a run to
    clone from."""
    para = cell.paragraphs[0]
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    if not text:
        return
    run = para.add_run(text)
    run.bold = bold
    run.font.name = 'Calibri'
    run.font.color.rgb = color


def set_column_widths(tbl_skeleton, colhdr_tr, content_tr, widths):
    grid = tbl_skeleton.find(qn('w:tblGrid'))
    for col, w in zip(grid.findall(qn('w:gridCol')), widths):
        col.set(qn('w:w'), str(w))
    for tr in (colhdr_tr, content_tr):
        for tc, w in zip(tr.findall(qn('w:tc')), widths):
            tc.find(qn('w:tcPr')).find(qn('w:tcW')).set(qn('w:w'), str(w))


def drop_first_two_columns(tbl_skeleton_4col, tr4, widths2):
    """Build a 2-column variant (gold accent + title only) of a 4-column
    skeleton/row — used for the poster tables, which don't need the
    Start/End time columns."""
    new_tbl = copy.deepcopy(tbl_skeleton_4col)
    grid = new_tbl.find(qn('w:tblGrid'))
    cols = grid.findall(qn('w:gridCol'))
    for col in cols[:2]:
        grid.remove(col)
    for col, w in zip(grid.findall(qn('w:gridCol')), widths2):
        col.set(qn('w:w'), str(w))

    new_tr = copy.deepcopy(tr4)
    tcs = new_tr.findall(qn('w:tc'))
    for tc in tcs[:2]:
        new_tr.remove(tc)
    for tc, w in zip(new_tr.findall(qn('w:tc')), widths2):
        tc.find(qn('w:tcPr')).find(qn('w:tcW')).set(qn('w:w'), str(w))

    return new_tbl, new_tr


def stamp_span(tr, span, width):
    """Deep-copy a single-cell caption/banner row and force its gridSpan
    and width to the given values, adding either element if the source
    row didn't already have one (e.g. a 1-column day-banner table's row,
    reused as a caption spanning a wider multi-column table)."""
    new_tr = copy.deepcopy(tr)
    tc = new_tr.find(qn('w:tc'))
    tcPr = tc.find(qn('w:tcPr'))

    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.insert(0, tcW)
    tcW.set(qn('w:type'), 'dxa')
    tcW.set(qn('w:w'), str(width))

    gridspan = tcPr.find(qn('w:gridSpan'))
    if gridspan is None:
        gridspan = OxmlElement('w:gridSpan')
        tcW.addnext(gridspan)
    gridspan.set(qn('w:val'), str(span))

    return new_tr


def set_cell_lines(cell, lines):
    """lines: list of (text, bold) tuples, one per paragraph in the cell.
    The cell's existing first paragraph (with its template run already on
    it) is reused for line 0; it's cloned for each additional line so the
    same font/color carries over."""
    base_para = cell.paragraphs[0]
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)

    if not lines:
        set_run(base_para, '')
        return

    text0, bold0 = lines[0]
    set_run(base_para, text0, bold0)

    for text, bold in lines[1:]:
        new_p_el = copy.deepcopy(base_para._p)
        cell._tc.append(new_p_el)
        new_p = Paragraph(new_p_el, cell)
        set_run(new_p, text, bold)


def table_from_trs(tbl_skeleton, trs):
    """Deep-copy a <w:tbl> skeleton (tblPr/tblGrid, no rows) and populate
    it with the given (already-cloned) <w:tr> elements."""
    new_tbl = copy.deepcopy(tbl_skeleton)
    for tr in new_tbl.findall(qn('w:tr')):
        new_tbl.remove(tr)
    for tr in trs:
        new_tbl.append(tr)
    return new_tbl


class Cursor:
    """Tracks the last-inserted body element so new paragraphs/tables can
    be spliced in immediately after it, in document order."""

    def __init__(self, doc, anchor_element):
        self.doc = doc
        self.el = anchor_element

    def insert(self, element):
        self.el.addnext(element)
        self.el = element

    def add_table(self, tbl_skeleton, trs):
        tbl_el = table_from_trs(tbl_skeleton, trs)
        self.insert(tbl_el)
        return Table(tbl_el, self.doc)

    def add_paragraph(self, lines, p_template):
        """lines: list of (text, bold) tuples rendered as separate
        paragraphs, styled like p_template (a <w:p> element carrying the
        font/color to clone)."""
        for text, bold in lines:
            p_el = copy.deepcopy(p_template)
            self.insert(p_el)
            p = Paragraph(p_el, self.doc)
            set_run(p, text, bold)

    def add_spacer(self, p_template):
        p_el = copy.deepcopy(p_template)
        self.insert(p_el)
        p = Paragraph(p_el, self.doc)
        for r in list(p.runs):
            r._element.getparent().remove(r._element)


# ---------------------------------------------------------------------------
# Timetable content
# ---------------------------------------------------------------------------

def build_schedule(cursor, rows, day_tbl_skel, colhdr_tr, content_tr, note_p_template):
    days = {}
    day_order = []
    for row in rows:
        date = row.get('date', '') or ''
        if date not in days:
            days[date] = []
            day_order.append(date)
        days[date].append(row)

    for date in day_order:
        day_rows = days[date]

        day_tr = stamp_span(day_tbl_skel['tr'], 4, TOTAL_WIDTH)
        day_tbl = cursor.add_table(day_tbl_skel['tbl'], [day_tr])
        set_run(day_tbl.rows[0].cells[0].paragraphs[0], fmt_day_banner(date), bold=None)
        cursor.add_spacer(note_p_template)

        current_session = None
        session_children = []

        def flush_session():
            nonlocal current_session
            if current_session is not None:
                title = current_session.get('title', '') or ''
                mods = current_session.get('speaker', '') or ''

                caption_tr = stamp_span(day_tbl_skel['tr'], 4, TOTAL_WIDTH)
                trs = [caption_tr, copy.deepcopy(colhdr_tr)]
                for child in session_children:
                    trs.append(copy.deepcopy(content_tr))
                sess_tbl = cursor.add_table(day_tbl_skel['tbl'], trs)

                cap_lines = [(title, True)]
                if mods:
                    cap_lines.append((mods, False))
                set_cell_lines(sess_tbl.rows[0].cells[0], cap_lines)
                set_styled_run(sess_tbl.rows[1].cells[2], '#', WHITE)

                for i, child in enumerate(session_children):
                    row = sess_tbl.rows[2 + i]
                    set_run(row.cells[0].paragraphs[0], fmt_time(child.get('start')))
                    set_run(row.cells[1].paragraphs[0], fmt_time(child.get('end')))

                    num = child.get('number', '') or ''
                    set_styled_run(row.cells[2], f'O-{num}' if num else '', NAVY)

                    ctitle = child.get('title', '') or ''
                    lines = [(ctitle, True)]
                    spk = speaker_line(child)
                    if spk:
                        lines.append((spk, False))
                    cnotes = child.get('notes', '') or ''
                    if cnotes:
                        lines.append((cnotes, False))
                    set_cell_lines(row.cells[3], lines)

                cursor.add_spacer(note_p_template)
            current_session = None
            session_children.clear()

        for row in day_rows:
            rtype = (row.get('type', '') or '').lower()
            if rtype == 'notes':
                continue
            if rtype in CHILD_TYPES:
                session_children.append(row)
            elif rtype == 'session':
                flush_session()
                current_session = row
            else:
                flush_session()
                title = row.get('title', '') or ''
                lines = [(f"{fmt_time(row.get('start'))}–{fmt_time(row.get('end'))}  {title}", True)]
                spk = speaker_line(row)
                if spk:
                    lines.append((spk, False))
                notes = row.get('notes', '') or ''
                if notes:
                    lines.append((notes, False))
                cursor.add_paragraph(lines, note_p_template)
                cursor.add_spacer(note_p_template)

        flush_session()


POSTER_COL_WIDTHS = (720, 10080)  # gold accent (#) + title/presenter


def build_posters(cursor, posters, day_tbl_skel, colhdr_tr, content_tr, note_p_template):
    if not posters:
        return

    categories = []
    by_category = {}
    for row in posters:
        cat = row.get('category', '') or 'Other'
        if cat not in by_category:
            by_category[cat] = []
            categories.append(cat)
        by_category[cat].append(row)

    banner_tr = stamp_span(day_tbl_skel['tr'], 4, TOTAL_WIDTH)
    banner_tbl = cursor.add_table(day_tbl_skel['tbl'], [banner_tr])
    set_run(banner_tbl.rows[0].cells[0].paragraphs[0], 'POSTER SESSION')
    cursor.add_spacer(note_p_template)

    cursor.add_paragraph([(
        f'All {len(posters)} poster-assigned abstracts are presented in the area adjacent to the main '
        'conference room during Session III — Poster Presentations (1:00 PM – 2:30 PM, October 8, 2026), '
        'grouped below by topic.', False,
    )], note_p_template)
    cursor.add_spacer(note_p_template)

    # Posters don't need Start/End time columns, so drop them and keep just
    # the gold accent (#) and title/presenter columns.
    poster_tbl_skel, poster_colhdr_tr = drop_first_two_columns(day_tbl_skel['tbl'], colhdr_tr, POSTER_COL_WIDTHS)
    _, poster_content_tr = drop_first_two_columns(day_tbl_skel['tbl'], content_tr, POSTER_COL_WIDTHS)
    poster_caption_tr = stamp_span(day_tbl_skel['tr'], 2, sum(POSTER_COL_WIDTHS))

    for cat in categories:
        group = by_category[cat]

        cat_tr = copy.deepcopy(poster_caption_tr)
        colhdr = copy.deepcopy(poster_colhdr_tr)
        trs = [cat_tr, colhdr] + [copy.deepcopy(poster_content_tr) for _ in group]
        cat_tbl = cursor.add_table(poster_tbl_skel, trs)

        set_run(cat_tbl.rows[0].cells[0].paragraphs[0], f'{cat} ({len(group)})', bold=True)
        set_styled_run(cat_tbl.rows[1].cells[0], '#', WHITE)
        set_run(cat_tbl.rows[1].cells[1].paragraphs[0], 'Title')

        for i, row in enumerate(group):
            trow = cat_tbl.rows[2 + i]
            num = row.get('number', '') or ''
            set_styled_run(trow.cells[0], f'P-{num}' if num else '', NAVY)

            title = row.get('title', '') or ''
            author = row.get('author', '') or ''
            inst = row.get('institution', '') or ''
            city = row.get('city', '') or ''
            country = row.get('country', '') or ''
            presenter = author
            location = ', '.join(filter(None, [inst, city, country]))
            if location:
                presenter += f', {location}'
            lines = [(title, True)]
            if presenter:
                lines.append((presenter, False))
            set_cell_lines(trow.cells[1], lines)

        cursor.add_spacer(note_p_template)


# ---------------------------------------------------------------------------
# Invited speakers
# ---------------------------------------------------------------------------

def rebuild_invited_table(invited_tbl, speakers):
    """Replace the 3-column "Invited Speakers" roster's rows in place,
    reusing its existing banner/header/content rows as style templates.
    Unlike the author index, this table's column count never changes, so
    there's no format-detection branch needed here."""
    tbl_el = invited_tbl._tbl
    trs = tbl_el.findall(qn('w:tr'))
    banner_tr, colhdr_tr, content_tr = (copy.deepcopy(tr) for tr in trs[:3])
    for tr in trs:
        tbl_el.remove(tr)

    tbl_el.append(banner_tr)
    tbl_el.append(colhdr_tr)

    tbl = Table(tbl_el, invited_tbl._parent)
    set_run(tbl.rows[0].cells[0].paragraphs[0], 'Invited Speakers')

    for speaker in speakers:
        trow_el = copy.deepcopy(content_tr)
        tbl_el.append(trow_el)
        row = tbl.rows[-1]
        for text, cell in zip(
            (speaker['name'], speaker['affiliation'], speaker['location']),
            row.cells,
        ):
            set_run(cell.paragraphs[0], text)


# ---------------------------------------------------------------------------
# Author index
# ---------------------------------------------------------------------------

def rebuild_author_table(author_tbl, authors):
    """Replace the roster table's rows in place with a 4-column author
    index (Name | Affiliation | City, Country | Presentations), reusing
    its existing banner/header/content rows as style templates so fonts,
    shading, and borders carry over unchanged.

    The source table may either be the original 3-column "Regular Session
    Speakers" table, or this script's own previously-generated 4-column
    "Author Index" (e.g. if that output was saved back over the source
    docx) — detected from the column-header row's cell count, so re-running
    against either shape works without adding a spurious 5th column."""
    tbl_el = author_tbl._tbl
    trs = tbl_el.findall(qn('w:tr'))
    banner_tr, colhdr_tr, content_tr = (copy.deepcopy(tr) for tr in trs[:3])
    for tr in trs:
        tbl_el.remove(tr)

    total_width = sum(AUTHOR_COL_WIDTHS)
    already_4col = len(colhdr_tr.findall(qn('w:tc'))) == 4

    # gridCol count: add a 4th (cloned from the last existing one) only if
    # the source table doesn't already have one; then set all widths.
    grid = tbl_el.find(qn('w:tblGrid'))
    if not already_4col:
        grid.append(copy.deepcopy(grid.findall(qn('w:gridCol'))[-1]))
    for col, w in zip(grid.findall(qn('w:gridCol')), AUTHOR_COL_WIDTHS):
        col.set(qn('w:w'), str(w))

    # Banner row: spans 3 columns -> 4 (or already 4), same total width,
    # new caption text.
    banner_tc = banner_tr.find(qn('w:tc'))
    banner_tcPr = banner_tc.find(qn('w:tcPr'))
    banner_tcPr.find(qn('w:tcW')).set(qn('w:w'), str(total_width))
    banner_tcPr.find(qn('w:gridSpan')).set(qn('w:val'), '4')

    # Column header + content rows: clone their last cell as the template
    # for the new 4th ("Presentations") column only if not already there,
    # then set all 4 widths.
    for tr in (colhdr_tr, content_tr):
        if not already_4col:
            tcs = tr.findall(qn('w:tc'))
            tr.append(copy.deepcopy(tcs[-1]))
        for tc, w in zip(tr.findall(qn('w:tc')), AUTHOR_COL_WIDTHS):
            tc.find(qn('w:tcPr')).find(qn('w:tcW')).set(qn('w:w'), str(w))

    tbl_el.append(banner_tr)
    tbl_el.append(colhdr_tr)

    tbl = Table(tbl_el, author_tbl._parent)
    set_run(tbl.rows[0].cells[0].paragraphs[0], 'Author Index')
    for text, cell in zip(('Name', 'Affiliation', 'City, Country', 'Presentations'), tbl.rows[1].cells):
        set_run(cell.paragraphs[0], text)

    for author in authors:
        trow_el = copy.deepcopy(content_tr)
        tbl_el.append(trow_el)
        row = tbl.rows[-1]
        for text, cell in zip(
            (author['name'], author['affiliation'], author['location'], author['presentations']),
            row.cells,
        ):
            set_run(cell.paragraphs[0], text)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def find_session_shaped_table(tables):
    """The first table with >=3 rows and a 4-cell second row — i.e. a real
    session table (caption + Start/End/#/Title header + >=1 talk), as
    opposed to a 1-row day/poster banner or a 2-column poster table."""
    for t in tables:
        if len(t.rows) >= 3 and len(t.rows[1]._tr.findall(qn('w:tc'))) == 4:
            return t
    raise RuntimeError('No 4-column session-shaped table found to use as a style template.')


def generate(src_path, out_path, program_csv, posters_csv, authors_csv):
    doc = docx.Document(src_path)
    tables = doc.tables

    # The source document may either be the original monolithic-table
    # booklet, or a previously-generated split-table version (each session
    # is already its own table). Detect which, and pull style templates
    # accordingly, so re-running this script against its own prior output
    # keeps working.
    # The two speaker-roster tables ("Invited Speakers", "Regular Session
    # Speakers"/"Author Index") are always the last two tables in the
    # document, in both the original monolithic format and this script's
    # own split-table output — so the schedule/posters to rebuild is
    # everything from the first table up to (not including) tables[-2].
    schedule_start_tbl = tables[0]._tbl
    schedule_end_tbl = tables[-2]._tbl
    invited_tbl = tables[-2]
    author_tbl = tables[-1]

    if len(tables[0].rows) > 1:
        # Original monolithic format: one big table, rows 1/2/4 are the
        # day banner / column header / content row templates.
        table = tables[0]
        day_tr = copy.deepcopy(table.rows[1]._tr)
        colhdr_tr = copy.deepcopy(table.rows[2]._tr)
        content_tr = copy.deepcopy(table.rows[4]._tr)
        tbl_skeleton = copy.deepcopy(table._tbl)
        for tr in tbl_skeleton.findall(qn('w:tr')):
            tbl_skeleton.remove(tr)
    else:
        # Already-split format: tables[0] is a 1-row day banner; find the
        # first real session table for the column-header/content templates.
        day_tr = copy.deepcopy(tables[0].rows[0]._tr)
        session_tpl = find_session_shaped_table(tables)
        colhdr_tr = copy.deepcopy(session_tpl.rows[1]._tr)
        content_tr = copy.deepcopy(session_tpl.rows[2]._tr)
        tbl_skeleton = copy.deepcopy(session_tpl._tbl)
        for tr in tbl_skeleton.findall(qn('w:tr')):
            tbl_skeleton.remove(tr)

    day_tbl_skel = {'tbl': tbl_skeleton, 'tr': day_tr}
    set_column_widths(tbl_skeleton, colhdr_tr, content_tr, COL_WIDTHS)

    # A plain-paragraph template (font/color) for standalone (non-table)
    # lines, cloned from the content row's title-column paragraph so it
    # matches the table text's look (Calibri, dark gray).
    note_p_template = copy.deepcopy(content_tr.findall(qn('w:tc'))[3].find(qn('w:p')))

    # Anchor: an empty paragraph inserted right before the schedule, used
    # as the starting point for splicing in the new content. Everything
    # from there up to (but not including) the first roster table — every
    # old day/session/poster table and every standalone paragraph between
    # them — is then removed.
    anchor_p = doc.add_paragraph()
    anchor_el = anchor_p._p
    anchor_el.getparent().remove(anchor_el)
    schedule_start_tbl.addprevious(anchor_el)

    body = doc.element.body
    children = list(body)
    start_idx = children.index(schedule_start_tbl)
    end_idx = children.index(schedule_end_tbl)
    for el in children[start_idx:end_idx]:
        el.getparent().remove(el)

    cursor = Cursor(doc, anchor_el)

    program_rows = load_csv(program_csv)
    build_schedule(cursor, program_rows, day_tbl_skel, colhdr_tr, content_tr, note_p_template)

    posters = load_csv(posters_csv) if Path(posters_csv).exists() else []
    build_posters(cursor, posters, day_tbl_skel, colhdr_tr, content_tr, note_p_template)

    invited_speakers = build_invited_speakers(program_rows)
    rebuild_invited_table(invited_tbl, invited_speakers)

    if Path(authors_csv).exists():
        subid_to_pres = build_subid_to_pres(program_rows, posters)
        authors = build_author_index(authors_csv, subid_to_pres)
        rebuild_author_table(author_tbl, authors)

    doc.save(out_path)
    print(f'Wrote: {out_path}')


if __name__ == '__main__':
    args = sys.argv[1:]
    src = args[0] if len(args) > 0 else str(SCRIPT_DIR / 'iMRI2026-program-booklet.docx')
    out = args[1] if len(args) > 1 else str(SCRIPT_DIR / 'iMRI2026-program-booklet-updated.docx')
    generate(
        src, out,
        SCRIPT_DIR / 'program.csv',
        SCRIPT_DIR / 'posters.csv',
        SCRIPT_DIR / 'authors_accepted.csv',
    )
